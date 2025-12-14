import os
import io
import json
from typing import Dict, List, Optional, Any
from datetime import datetime
import pandas as pd
import PyPDF2
import pdfplumber
from docx import Document as DocxDocument
from PIL import Image
import re

class DocumentProcessor:
    """Process various document types and extract text/data - PRODUCTION FIXED VERSION"""
    
    def __init__(self):
        self.supported_types = {
            'pdf': ['.pdf'],
            'excel': ['.xlsx', '.xls', '.csv'],
            'word': ['.docx', '.doc'],
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp'],
            'text': ['.txt', '.md']
        }
    
    def get_file_type(self, filename: str) -> Optional[str]:
        """Determine file type from extension"""
        ext = os.path.splitext(filename)[1].lower()
        for file_type, extensions in self.supported_types.items():
            if ext in extensions:
                return file_type
        return None
    
    def process_document(self, file_path: str, file_type: str = None) -> Dict[str, Any]:
        """Main entry point for document processing"""
        if not os.path.exists(file_path):
            return {'error': 'File not found', 'success': False}
        
        if not file_type:
            file_type = self.get_file_type(file_path)
        
        if not file_type:
            return {'error': 'Unsupported file type', 'success': False}
        
        try:
            if file_type == 'pdf':
                return self.process_pdf(file_path)
            elif file_type == 'excel':
                return self.process_excel(file_path)
            elif file_type == 'word':
                return self.process_word(file_path)
            elif file_type == 'image':
                return self.process_image(file_path)
            elif file_type == 'text':
                return self.process_text(file_path)
            else:
                return {'error': 'Unsupported file type', 'success': False}
        except Exception as e:
            return {'error': str(e), 'success': False}
    
    def process_pdf(self, file_path: str) -> Dict[str, Any]:
        """Extract text from PDF"""
        text_content = []
        metadata = {
            'num_pages': 0,
            'tables': [],
            'has_images': False
        }
        
        try:
            # Try pdfplumber first (better for tables)
            with pdfplumber.open(file_path) as pdf:
                metadata['num_pages'] = len(pdf.pages)
                
                for page_num, page in enumerate(pdf.pages, 1):
                    # Extract text
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append({
                            'page': page_num,
                            'content': page_text
                        })
                    
                    # Extract tables
                    tables = page.extract_tables()
                    if tables:
                        for table_num, table in enumerate(tables, 1):
                            metadata['tables'].append({
                                'page': page_num,
                                'table_num': table_num,
                                'data': table
                            })
                    
                    # Check for images
                    if page.images:
                        metadata['has_images'] = True
        
        except Exception as e:
            # Fallback to PyPDF2
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    metadata['num_pages'] = len(pdf_reader.pages)
                    
                    for page_num, page in enumerate(pdf_reader.pages, 1):
                        page_text = page.extract_text()
                        if page_text:
                            text_content.append({
                                'page': page_num,
                                'content': page_text
                            })
            except Exception as fallback_error:
                return {
                    'error': f'PDF processing failed: {str(e)}, Fallback error: {str(fallback_error)}',
                    'success': False
                }
        
        # Combine all text
        full_text = '\n\n'.join([page['content'] for page in text_content])
        
        return {
            'success': True,
            'file_type': 'pdf',
            'text_content': text_content,
            'metadata': metadata,
            'full_text': full_text
        }
    
    def process_excel(self, file_path: str) -> Dict[str, Any]:
        """Extract data from Excel/CSV files"""
        try:
            # Determine if it's CSV or Excel
            file_ext = os.path.splitext(file_path)[1].lower()
            
            sheets = {}
            
            if file_ext == '.csv':
                # Process CSV
                df = pd.read_csv(file_path)
                
                # Convert to records
                records = df.to_dict('records')
                
                sheets['Sheet1'] = {
                    'data': records,
                    'dataframe': df,
                    'columns': df.columns.tolist(),
                    'num_rows': len(df),
                    'num_cols': len(df.columns)
                }
            else:
                # Process Excel
                excel_file = pd.ExcelFile(file_path)
                
                for sheet_name in excel_file.sheet_names:
                    df = pd.read_excel(file_path, sheet_name=sheet_name)
                    
                    # Convert to records
                    records = df.to_dict('records')
                    
                    sheets[sheet_name] = {
                        'data': records,
                        'dataframe': df,
                        'columns': df.columns.tolist(),
                        'num_rows': len(df),
                        'num_cols': len(df.columns)
                    }
            
            # Create full text representation
            full_text_parts = []
            for sheet_name, sheet_data in sheets.items():
                full_text_parts.append(f"Sheet: {sheet_name}")
                full_text_parts.append(sheet_data['dataframe'].to_string())
            
            full_text = '\n\n'.join(full_text_parts)
            
            return {
                'success': True,
                'file_type': 'excel',
                'sheets': sheets,
                'metadata': {
                    'num_sheets': len(sheets),
                    'sheet_names': list(sheets.keys())
                },
                'full_text': full_text
            }
        
        except Exception as e:
            return {'error': f'Excel processing failed: {str(e)}', 'success': False}
    
    def process_word(self, file_path: str) -> Dict[str, Any]:
        """Extract text from Word documents"""
        try:
            doc = DocxDocument(file_path)
            
            # Extract paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Extract tables
            tables = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables.append(table_data)
            
            # Combine all text
            full_text = '\n\n'.join(paragraphs)
            
            return {
                'success': True,
                'file_type': 'word',
                'paragraphs': paragraphs,
                'tables': tables,
                'metadata': {
                    'num_paragraphs': len(paragraphs),
                    'num_tables': len(tables)
                },
                'full_text': full_text
            }
        
        except Exception as e:
            return {'error': f'Word processing failed: {str(e)}', 'success': False}
    
    def process_image(self, file_path: str) -> Dict[str, Any]:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            
            # Perform OCR
            try:
                import pytesseract
                text = pytesseract.image_to_string(image)
            except Exception as ocr_error:
                text = ""
                print(f"OCR failed: {ocr_error}. Install Tesseract OCR for image text extraction.")
            
            metadata = {
                'width': image.width,
                'height': image.height,
                'format': image.format,
                'mode': image.mode
            }
            
            return {
                'success': True,
                'file_type': 'image',
                'text_content': text,
                'metadata': metadata,
                'full_text': text
            }
        
        except Exception as e:
            return {'error': f'Image processing failed: {str(e)}', 'success': False}
    
    def process_text(self, file_path: str) -> Dict[str, Any]:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            
            lines = text.split('\n')
            
            return {
                'success': True,
                'file_type': 'text',
                'text_content': text,
                'metadata': {
                    'num_lines': len(lines),
                    'num_chars': len(text)
                },
                'full_text': text
            }
        
        except Exception as e:
            return {'error': f'Text processing failed: {str(e)}', 'success': False}
    
    def extract_financial_data(self, processed_doc: Dict[str, Any], doc_type: str) -> List[Dict[str, Any]]:
        """Extract structured financial data based on document type - FIXED VERSION"""
        financial_data = []
        
        try:
            if doc_type == 'salary_slip':
                financial_data = self._extract_salary_data(processed_doc)
            elif doc_type == 'bank_statement':
                financial_data = self._extract_bank_transactions(processed_doc)
            elif doc_type == 'credit_card_statement':
                financial_data = self._extract_credit_card_data(processed_doc)
            elif doc_type == 'loan_data':
                financial_data = self._extract_loan_data(processed_doc)
            elif doc_type == 'investment_statement' or doc_type == 'collateral_details':
                financial_data = self._extract_investment_data(processed_doc)
        except Exception as e:
            print(f"Error extracting financial data: {e}")
            import traceback
            traceback.print_exc()
        
        return financial_data
    
    def _extract_numbers_from_text(self, text: str) -> List[float]:
        """Extract all numbers from text"""
        pattern = r'₹?\s*([0-9,]+\.?[0-9]*)'
        matches = re.findall(pattern, text)
        numbers = []
        for match in matches:
            try:
                num = float(match.replace(',', ''))
                numbers.append(num)
            except:
                continue
        return numbers
    
    def _extract_salary_data(self, doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract salary information - FIXED VERSION"""
        data = []
        text = doc.get('full_text', '')
        
        # Patterns for salary components
        patterns = {
            'CTC': [
                r'Annual\s+CTC[:\s■]+₹?\s*([\d,]+)',
                r'CTC[:\s■]+₹?\s*([\d,]+)',
            ],
            'Gross Salary': [
                r'Gross\s+Salary[:\s]+₹?\s*([\d,]+)',
                r'Annual.*?Gross.*?₹?\s*([\d,]+)',
            ],
            'Basic Salary': [
                r'Basic\s+Salary[:\s]+₹?\s*([\d,]+)',
                r'Basic[:\s]+₹?\s*([\d,]+)',
            ],
            'Net Salary': [
                r'Net.*?Salary[:\s]+₹?\s*([\d,]+)',
                r'Take\s+Home[:\s]+₹?\s*([\d,]+)',
            ]
        }
        
        extracted_amounts = {}
        
        for category, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    amount_str = match.group(1).replace(',', '')
                    try:
                        amount = float(amount_str)
                        
                        # Validation: Annual amounts should be > 100k
                        if category in ['CTC', 'Gross Salary', 'Basic Salary']:
                            if amount < 100000:
                                # Likely monthly, multiply by 12
                                amount = amount * 12
                        
                        extracted_amounts[category] = amount
                        break
                    except:
                        continue
        
        # Validate: CTC should be highest
        if extracted_amounts:
            max_amount = max(extracted_amounts.values())
            if 'CTC' in extracted_amounts and extracted_amounts['CTC'] < max_amount:
                extracted_amounts['CTC'] = max_amount
        
        # Convert to data entries
        for category, amount in extracted_amounts.items():
            if amount > 0:
                data.append({
                    'type': 'income',
                    'category': category,
                    'amount': amount,
                    'description': f'{category} from salary slip',
                    'extracted_from': 'salary_slip'
                })
        
        return data
    
    def _extract_bank_transactions(self, doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract bank transactions - FIXED VERSION"""
        data = []
        
        # For Excel/CSV files
        if doc.get('file_type') == 'excel':
            sheets = doc.get('sheets', {})
            
            for sheet_name, sheet_data in sheets.items():
                df = sheet_data.get('dataframe')
                
                if df is None or df.empty:
                    continue
                
                # Find column names
                date_col = None
                desc_col = None
                debit_col = None
                credit_col = None
                amount_col = None
                
                for col in df.columns:
                    col_lower = str(col).lower()
                    
                    if any(x in col_lower for x in ['date', 'txn date', 'trans date', 'transaction date']):
                        date_col = col
                    elif any(x in col_lower for x in ['description', 'narration', 'particulars', 'details']):
                        desc_col = col
                    elif any(x in col_lower for x in ['debit', 'withdrawal', 'dr', 'debit amount']):
                        debit_col = col
                    elif any(x in col_lower for x in ['credit', 'deposit', 'cr', 'credit amount']):
                        credit_col = col
                    elif 'amount' in col_lower and not any(x in col_lower for x in ['debit', 'credit']):
                        amount_col = col
                
                # Process each row
                for idx, row in df.iterrows():
                    try:
                        # Get amount
                        amount = 0
                        trans_type = 'expense'
                        
                        if debit_col and pd.notna(row[debit_col]):
                            amount_str = str(row[debit_col]).replace(',', '').replace('₹', '').strip()
                            try:
                                amount = float(amount_str)
                                trans_type = 'expense'
                            except:
                                continue
                        elif credit_col and pd.notna(row[credit_col]):
                            amount_str = str(row[credit_col]).replace(',', '').replace('₹', '').strip()
                            try:
                                amount = float(amount_str)
                                trans_type = 'income'
                            except:
                                continue
                        elif amount_col and pd.notna(row[amount_col]):
                            amount_str = str(row[amount_col]).replace(',', '').replace('₹', '').strip()
                            try:
                                amount = abs(float(amount_str))
                                trans_type = 'expense'
                            except:
                                continue
                        
                        # CRITICAL: Validate amount is reasonable
                        if amount <= 0 or amount > 10000000:  # Skip if > 1 crore
                            continue
                        
                        # Get description
                        description = "Bank Transaction"
                        if desc_col and pd.notna(row[desc_col]):
                            description = str(row[desc_col])[:200]
                        
                        # Get date
                        trans_date = None
                        if date_col and pd.notna(row[date_col]):
                            try:
                                trans_date = pd.to_datetime(row[date_col])
                            except:
                                pass
                        
                        # Categorize based on description
                        category = self._categorize_transaction(description)
                        
                        data.append({
                            'type': trans_type,
                            'category': category,
                            'amount': amount,
                            'description': description,
                            'transaction_date': trans_date,
                            'extracted_from': 'bank_statement'
                        })
                    
                    except Exception as e:
                        print(f"Error processing row {idx}: {e}")
                        continue
        
        # For PDF files
        else:
            text = doc.get('full_text', '')
            
            # Try to extract from tables first
            tables = doc.get('metadata', {}).get('tables', [])
            
            if tables:
                # Process tables
                for table_info in tables:
                    table_data = table_info.get('data', [])
                    
                    if not table_data or len(table_data) < 2:
                        continue
                    
                    # Try to convert table to DataFrame
                    try:
                        # Use first row as headers
                        headers = table_data[0]
                        rows = table_data[1:]
                        
                        df = pd.DataFrame(rows, columns=headers)
                        
                        # Process similar to Excel (same logic as above)
                        # Find columns
                        date_col = None
                        desc_col = None
                        debit_col = None
                        credit_col = None
                        
                        for col in df.columns:
                            col_lower = str(col).lower()
                            
                            if any(x in col_lower for x in ['date', 'txn', 'trans']):
                                date_col = col
                            elif any(x in col_lower for x in ['description', 'narration', 'particulars']):
                                desc_col = col
                            elif any(x in col_lower for x in ['debit', 'withdrawal', 'dr']):
                                debit_col = col
                            elif any(x in col_lower for x in ['credit', 'deposit', 'cr']):
                                credit_col = col
                        
                        # Process rows
                        for idx, row in df.iterrows():
                            try:
                                amount = 0
                                trans_type = 'expense'
                                
                                if debit_col and pd.notna(row[debit_col]):
                                    amount_str = str(row[debit_col]).replace(',', '').replace('₹', '').strip()
                                    try:
                                        amount = float(amount_str)
                                        trans_type = 'expense'
                                    except:
                                        continue
                                elif credit_col and pd.notna(row[credit_col]):
                                    amount_str = str(row[credit_col]).replace(',', '').replace('₹', '').strip()
                                    try:
                                        amount = float(amount_str)
                                        trans_type = 'income'
                                    except:
                                        continue
                                
                                if amount <= 0 or amount > 10000000:
                                    continue
                                
                                description = "Bank Transaction"
                                if desc_col and pd.notna(row[desc_col]):
                                    description = str(row[desc_col])[:200]
                                
                                category = self._categorize_transaction(description)
                                
                                data.append({
                                    'type': trans_type,
                                    'category': category,
                                    'amount': amount,
                                    'description': description,
                                    'extracted_from': 'bank_statement'
                                })
                            except:
                                continue
                    except:
                        continue
            
            # Fallback: Try regex extraction from text
            if not data:
                lines = text.split('\n')
                
                date_pattern = r'(\d{2}[/-]\d{2}[/-]\d{4})'
                amount_pattern = r'₹?\s*([\d,]+\.?\d*)'
                
                for line in lines:
                    # Skip header lines
                    if any(h in line.lower() for h in ['date', 'transaction', 'description', 'balance', 'particulars']):
                        continue
                    
                    # Look for date
                    date_match = re.search(date_pattern, line)
                    
                    if date_match:
                        # Extract amounts
                        amounts = re.findall(amount_pattern, line)
                        
                        if amounts:
                            parsed_amounts = []
                            for amt_str in amounts:
                                try:
                                    amt = float(amt_str.replace(',', ''))
                                    if 0 < amt < 10000000:  # Reasonable range
                                        parsed_amounts.append(amt)
                                except:
                                    continue
                            
                            if parsed_amounts:
                                # Usually first amount is transaction amount
                                amount = parsed_amounts[0]
                                
                                # Get description (text between date and amount)
                                desc_start = date_match.end()
                                try:
                                    desc_end = line.find(amounts[0], desc_start)
                                    description = line[desc_start:desc_end].strip()[:200] or "Transaction"
                                except:
                                    description = "Transaction"
                                
                                category = self._categorize_transaction(description)
                                
                                data.append({
                                    'type': 'expense',  # Default to expense
                                    'category': category,
                                    'amount': amount,
                                    'description': description,
                                    'extracted_from': 'bank_statement'
                                })
        
        return data
    
    def _categorize_transaction(self, description: str) -> str:
        """Categorize transaction based on description"""
        desc_lower = description.lower()
        
        categories = {
            'Shopping': ['amazon', 'flipkart', 'myntra', 'shop', 'mall', 'purchase'],
            'Food & Dining': ['zomato', 'swiggy', 'restaurant', 'food', 'cafe', 'dining'],
            'Utilities': ['electricity', 'water', 'gas', 'bill', 'bsnl', 'airtel', 'utility'],
            'UPI': ['upi', 'paytm', 'phonepe', 'gpay', 'googlepay'],
            'EMI': ['emi', 'loan'],
            'ATM Withdrawal': ['atm', 'withdrawal', 'cash'],
            'Transfer': ['transfer', 'neft', 'rtgs', 'imps'],
            'Salary': ['salary', 'wages', 'income'],
        }
        
        for category, keywords in categories.items():
            if any(kw in desc_lower for kw in keywords):
                return category
        
        return 'Bank Transaction'
    
    def _extract_credit_card_data(self, doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract credit card data - ENHANCED WITH TRANSACTION EXTRACTION"""
        text = doc.get('full_text', '')
        data = []
        
        # Detect bank type
        is_axis = 'axis bank' in text.lower()
        is_hdfc = 'hdfc bank' in text.lower()
        
        # Extract summary fields first
        summary_patterns = {
            'Total Amount Due': [
                r'Total\s+(?:Payment\s+)?Due[:\s]+(?:₹|C|Rs\.?)?\s*([\d,]+\.?\d*)',
                r'TOTAL\s+AMOUNT\s+DUE[:\s]+(?:₹|C)?\s*([\d,]+\.?\d*)',
            ],
            'Minimum Payment': [
                r'Minimum\s+(?:Payment\s+)?Due[:\s]+(?:₹|C|Rs\.?)?\s*([\d,]+\.?\d*)',
                r'MINIMUM\s+DUE[:\s]+(?:₹|C)?\s*([\d,]+\.?\d*)',
            ],
            'Credit Limit': [
                r'Credit\s+Limit[:\s]+(?:₹|C|Rs\.?)?\s*([\d,]+\.?\d*)',
                r'TOTAL\s+CREDIT\s+LIMIT[:\s]+(?:₹|C)?\s*([\d,]+\.?\d*)',
            ],
        }
        
        for category, pattern_list in summary_patterns.items():
            for pattern in pattern_list:
                matches = re.findall(pattern, text, re.IGNORECASE)
                if matches:
                    try:
                        amount = float(matches[0].replace(',', ''))
                        if 0 < amount < 100000000:
                            data.append({
                                'type': 'debt',
                                'category': category,
                                'amount': amount,
                                'description': f'{category} from credit card statement',
                                'extracted_from': 'credit_card_statement'
                            })
                            break
                    except:
                        continue
        
        # Extract individual transactions
        transactions = []
        
        if is_axis:
            # Axis Bank format: DATE | TRANSACTION DETAILS | CATEGORY | AMOUNT | CASHBACK
            lines = text.split('\n')
            for line in lines:
                date_match = re.match(r'(\d{2}/\d{2}/\d{4})\s+(.+)', line)
                if date_match:
                    date_str = date_match.group(1)
                    rest = date_match.group(2)
                    
                    amounts = re.findall(r'([\d,]+\.?\d*)\s+(Dr|Cr)', rest)
                    if amounts:
                        try:
                            amount = float(amounts[0][0].replace(',', ''))
                            amount_type = amounts[0][1]
                            
                            if 0 < amount < 1000000:
                                desc_end = rest.find(amounts[0][0])
                                description = rest[:desc_end].strip()[:200]
                                trans_type = 'expense' if amount_type == 'Dr' else 'income'
                                
                                try:
                                    trans_date = pd.to_datetime(date_str, format='%d/%m/%Y')
                                except:
                                    trans_date = None
                                
                                transactions.append({
                                    'type': trans_type,
                                    'category': self._categorize_transaction(description),
                                    'amount': amount,
                                    'description': description,
                                    'transaction_date': trans_date,
                                    'extracted_from': 'credit_card_statement'
                                })
                        except:
                            continue
        
        elif is_hdfc:
            # HDFC format: DATE & TIME | TRANSACTION DESCRIPTION | REWARDS | AMOUNT
            lines = text.split('\n')
            for line in lines:
                date_match = re.match(r'(\d{2}/\d{2}/\d{4})\|\s*(\d{2}:\d{2})\s+(.+)', line)
                if date_match:
                    date_str = date_match.group(1)
                    rest = date_match.group(3)
                    
                    amounts = re.findall(r'[C₹]\s*([\d,]+\.?\d*)', rest)
                    if amounts:
                        try:
                            amount = float(amounts[-1].replace(',', ''))
                            
                            if 0 < amount < 10000000:
                                desc_end = rest.rfind(amounts[-1])
                                description = rest[:desc_end].strip()
                                description = re.sub(r'\+\s*\d+\s*$', '', description).strip()[:200]
                                
                                trans_type = 'income' if 'payment' in description.lower() else 'expense'
                                
                                try:
                                    trans_date = pd.to_datetime(date_str, format='%d/%m/%Y')
                                except:
                                    trans_date = None
                                
                                transactions.append({
                                    'type': trans_type,
                                    'category': self._categorize_transaction(description),
                                    'amount': amount,
                                    'description': description,
                                    'transaction_date': trans_date,
                                    'extracted_from': 'credit_card_statement'
                                })
                        except:
                            continue
        
        data.extend(transactions)
        print(f"Credit card extraction: {len(data)} total ({len(data)-len(transactions)} summary + {len(transactions)} transactions)")
        
        return data
    
    def _extract_loan_data(self, doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract loan information - FIXED VERSION"""
        data = []
        text = doc.get('full_text', '')
        
        patterns = {
            'Principal Amount': [
                r'Principal[:\s]+₹?\s*([\d,]+\.?\d*)',
                r'Loan\s+Amount[:\s]+₹?\s*([\d,]+\.?\d*)',
                r'Outstanding\s+Principal[:\s]+₹?\s*([\d,]+\.?\d*)',
            ],
            'EMI': [
                r'EMI[:\s]+₹?\s*([\d,]+\.?\d*)',
                r'Monthly\s+Installment[:\s]+₹?\s*([\d,]+\.?\d*)',
                r'Monthly\s+EMI[:\s]+₹?\s*([\d,]+\.?\d*)',
            ],
            'Interest Rate': [
                r'Interest\s+Rate[:\s]+([\d.]+)%',
                r'ROI[:\s]+([\d.]+)%',
                r'Rate\s+of\s+Interest[:\s]+([\d.]+)%',
            ]
        }
        
        for category, pattern_list in patterns.items():
            for pattern in pattern_list:
                match = re.search(pattern, text, re.IGNORECASE)
                if match:
                    if category == 'Interest Rate':
                        try:
                            amount = float(match.group(1))
                        except:
                            continue
                    else:
                        amount_str = match.group(1).replace(',', '')
                        try:
                            amount = float(amount_str)
                        except:
                            continue
                    
                    if amount > 0:
                        data.append({
                            'type': 'debt',
                            'category': category,
                            'amount': amount,
                            'description': f'{category} from loan document',
                            'extracted_from': 'loan_data'
                        })
                    break
        
        return data
    
    def _extract_investment_data(self, doc: Dict[str, Any]) -> List[Dict[str, Any]]:
        """Extract investment information from holdings/portfolio"""
        data = []
        
        # For Excel/CSV files
        if doc.get('file_type') == 'excel':
            sheets = doc.get('sheets', {})
            
            for sheet_name, sheet_data in sheets.items():
                df = sheet_data.get('dataframe')
                
                if df is None or df.empty:
                    continue
                
                # Find column names
                name_col = None
                value_col = None
                qty_col = None
                price_col = None
                
                for col in df.columns:
                    col_lower = str(col).lower()
                    
                    if any(x in col_lower for x in ['symbol', 'scrip', 'instrument', 'name', 'company']):
                        name_col = col
                    elif any(x in col_lower for x in ['current value', 'market value', 'value']):
                        value_col = col
                    elif any(x in col_lower for x in ['quantity', 'qty', 'units']):
                        qty_col = col
                    elif any(x in col_lower for x in ['current price', 'ltp', 'price', 'nav']):
                        price_col = col
                
                # Process each row
                for idx, row in df.iterrows():
                    try:
                        # Skip rows with all NaN
                        if row.isna().all():
                            continue
                        
                        # Get investment name
                        name = None
                        if name_col and pd.notna(row[name_col]):
                            name = str(row[name_col]).strip()
                        
                        if not name or name == 'nan' or len(name) < 2:
                            continue
                        
                        # Get value
                        value = 0
                        
                        if value_col and pd.notna(row[value_col]):
                            try:
                                value = float(str(row[value_col]).replace(',', ''))
                            except:
                                pass
                        
                        # Calculate value from quantity and price if value not available
                        if value == 0 and qty_col and price_col:
                            try:
                                qty = float(str(row[qty_col]).replace(',', ''))
                                price = float(str(row[price_col]).replace(',', ''))
                                value = qty * price
                            except:
                                pass
                        
                        # Skip if no valid value or value too high
                        if value <= 0 or value > 1e10:
                            continue
                        
                        # Determine category based on sheet name or investment type
                        category = 'Investment'
                        if 'equity' in sheet_name.lower() or 'stock' in sheet_name.lower():
                            category = 'Equity'
                        elif 'mutual' in sheet_name.lower() or 'mf' in sheet_name.lower():
                            category = 'Mutual Funds'
                        
                        data.append({
                            'type': 'investment',
                            'category': category,
                            'amount': value,
                            'description': f'{name} - Current Value',
                            'extracted_from': 'investment_statement'
                        })
                    
                    except Exception as e:
                        print(f"Error processing investment row {idx}: {e}")
                        continue
        
        return data
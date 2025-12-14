import os
import re
import config
import logging
from docx import Document

logger = logging.getLogger(__name__)

class FileSaverAgent:
    def _sanitize_filename(self, title: str) -> str:
        """Creates a safe filename from a title."""
        # Remove special chars, spaces to underscores
        # Keep alphanumeric, spaces, hyphens
        s = re.sub(r'[^a-zA-Z0-9\s-]', '', title)
        s = s.strip().replace(' ', '_')
        return s[:50] # Limit length

    def save_as_txt(self, content: str, title: str) -> str:
        """Saves content to a .txt file."""
        try:
            filename = f"{self._sanitize_filename(title)}.txt"
            filepath = os.path.join(config.OUTPUT_DIR, filename)
            
            logger.info(f"Saving TXT draft to: {filepath}")
            
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(content)
            
            return filepath
        except Exception as e:
            logger.error(f"Failed to save .txt file: {e}")
            raise

    def save_as_docx(self, content: str, title: str) -> str:
        """Saves content to a .docx file."""
        try:
            filename = f"{self._sanitize_filename(title)}.docx"
            filepath = os.path.join(config.OUTPUT_DIR, filename)
            
            logger.info(f"Saving DOCX draft to: {filepath}")
            
            doc = Document()
            doc.add_heading(title, 0)
            
            # Simple parsing of markdown-like text to docx paragraphs
            # This is a basic conversion; it won't handle complex markdown perfectly
            for line in content.split('\n'):
                line = line.strip()
                if not line:
                    continue
                if line.startswith('# '):
                   doc.add_heading(line[2:], level=1)
                elif line.startswith('## '):
                   doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                   doc.add_heading(line[4:], level=3)
                elif line.startswith('- ') or line.startswith('* '):
                    doc.add_paragraph(line[2:], style='List Bullet')
                else:
                    doc.add_paragraph(line)
            
            doc.save(filepath)
            return filepath
        except Exception as e:
            logger.error(f"Failed to save .docx file: {e}")
            # Non-critical, just log it
            return ""
